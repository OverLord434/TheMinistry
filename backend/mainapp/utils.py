import io
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from django.core.mail import EmailMessage
from django.conf import settings
from datetime import datetime
from .config import sections

SECTION_PATHS = {s['name']: s['path'] for s in sections}

# Placeholder excerpt from Rosobrnadzor guidelines.
DOC_EXCERPT = (
    "По данному разделу обнаружены отсутствующие атрибуты. "
    "Пожалуйста ознакомьтесь с полной версией документа по ссылке: "
    "https://obrnadzor.gov.ru/wp-content/uploads/2024/09/"
    "metodicheskie-rekomendaczii-predstavleniya-informaczii-ob-"
    "obrazovatelnoj-or....pdf"
)


def add_hyperlink(paragraph, url: str, text: str | None = None):
    """Add a clickable hyperlink to a paragraph."""
    text = text or url
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def _split_top_level(s: str) -> list[str]:
    """
    Разбивает строку по запятым на верхнем уровне вложенности (не внутри скобок).
    """
    items = []
    buf = []
    depth = 0
    for ch in s:
        if ch == '(':
            depth += 1
        elif ch == ')':
            depth = max(depth - 1, 0)
        # если топ-уровневая запятая — закрываем текущий элемент
        if ch == ',' and depth == 0:
            items.append(''.join(buf).strip())
            buf = []
        else:
            buf.append(ch)
    last = ''.join(buf).strip()
    if last:
        items.append(last)
    return items


def format_attributes(attr_str: str) -> tuple[str, int]:
    """
    Форматирует строку атрибутов:
    - убирает переносы строк и дефисы переноса;
    - разделяет только по «верхнеуровневым» запятым;
    - формирует маркированный список и считает элементы.
    """
    if not attr_str:
        return "Нет данных", 0

    # 1) Сводим воедино разорванные переносами куски:
    cleaned = re.sub(r'\n-?\s*', ' ', attr_str)
    # 2) Убираем лишние пробелы подряд:
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()

    # 3) Разбиваем по запятым вне скобок:
    raw_items = _split_top_level(cleaned)

    # 4) Формируем вывод и считаем:
    bullets = []
    for itm in raw_items:
        if '(' in itm and ')' in itm:
            name, desc = itm.split('(', 1)
            desc = desc.rstrip(')')
            bullets.append(f"• {name.strip()}: {desc.strip()}")
        else:
            bullets.append(f"• {itm.strip()}")

    return '\n'.join(bullets), len(raw_items)


def generate_word(results):
    document = Document()
    
    # Настройка стилей
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Заголовок отчёта
    heading = document.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_run = heading.add_run('ОТЧЁТ ПРОВЕРКИ САЙТА\n')
    heading_run.font.size = Pt(14)
    heading_run.bold = True

    base_url = results[0].get('Адрес сайта', '') if results else ''
    if base_url:
        link_para = document.add_paragraph()
        link_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        link_para.add_run('Сайт организации: ')
        add_hyperlink(link_para, base_url, base_url)
        document.add_paragraph('')

    # Сбор общей статистики
    total_found = 0
    total_missing = 0
    for idx, item in enumerate(results):
        # Пропускаем дубли основного сайта
        if item['Раздел сайта'] == 'Основной сайт' and idx != 0:
            continue
        found_str = item.get('Найдено атрибутов', '')
        missing_str = item.get('Отсутствуют', '')
        _, found = format_attributes(found_str)
        _, missing = format_attributes(missing_str)
        total_found += found
        total_missing += missing

    # Добавляем блок общей статистики
    stats = document.add_paragraph()
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    stats.add_run("📊 Общая статистика\n").bold = True
    stats.add_run(f"✅ Найдено атрибутов: {total_found}\n❌ Отсутствует атрибутов: {total_missing}\n\n")
    # Разделитель
    document.add_paragraph('―' * 50).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Основное содержимое
    for idx, item in enumerate(results):
        if item['Раздел сайта'] == 'Основной сайт' and idx != 0:
            continue

        # Форматируем атрибуты
        found_formatted, _ = format_attributes(item.get('Найдено атрибутов', ''))
        missing_formatted, _ = format_attributes(item.get('Отсутствуют', ''))

        # Заголовок раздела
        section_header = document.add_paragraph()
        section_header.add_run(f"🔍 {item['Раздел сайта']}").bold = True
        section_header.add_run(f" [Статус: {item['Статус']}]")
        link_line = document.add_paragraph()
        url = item.get('Адрес сайта')
        if url:
            link_line.add_run('Ссылка: ')
            add_hyperlink(link_line, url, url)
        else:
            expected = base_url.rstrip('/') + SECTION_PATHS.get(item['Раздел сайта'], '')
            warn_run = link_line.add_run(
                f"Возможно у вас есть данный раздел, но он должен находиться по ссылке: {expected}"
            )
            warn_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Блок с атрибутами
        document.add_paragraph().add_run("Найдено:").bold = True
        document.add_paragraph(found_formatted)
        
        document.add_paragraph().add_run("Отсутствуют:").bold = True
        document.add_paragraph(missing_formatted)
        
        if 'Нет данных' not in missing_formatted:
            snippet = document.add_paragraph()
            run = snippet.add_run(DOC_EXCERPT)
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Разделитель между секциями
        if idx < len(results) - 1:
            document.add_paragraph('―' * 50).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


def send_word_report(to_email, org_name, results, subject=None, body=None):
    subject = subject or f"📄 Отчёт проверки для {org_name}"
    body = body or "Во вложении находится подробный отчёт по результатам проверки сайта."
    
    output = generate_word(results)
    if not output:
        return

    email = EmailMessage(
        subject,
        body,
        settings.DEFAULT_FROM_EMAIL,
        [to_email]
    )
    filename = f"Отчёт_{org_name}_{datetime.now().strftime('%d-%m-%Y_%H-%M')}.docx"
    email.attach(
        filename,
        output.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    try:
        email.send()
    except Exception as e:
        print(f"Ошибка отправки письма: {e}")
